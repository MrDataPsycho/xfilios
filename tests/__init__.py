from docx import Document


def get_records():
    records = (
        (3, "101", "Spam"),
        (7, "422", "Eggs"),
        (4, "631", "Spam, spam, eggs, and spam"),
    )
    return records


def create_demo_docx():
    docx = Document()
    docx.add_heading("Document Title", 0)
    p = docx.add_paragraph("A plain paragraph having some")
    p.add_run("bold").bold = True
    records = get_records()
    table = docx.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Id"
    hdr_cells[2].text = "Desc"
    for qty, _id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = _id
        row_cells[2].text = desc
    docx.add_page_break()
    docx.add_paragraph("A plain paragraph having some New")
    return docx
